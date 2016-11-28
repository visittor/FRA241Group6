#!/usr/bin/env python
# -*- coding: utf-8 -*-

import os
import sys
import transaction
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor

from pyramid.paster import (
    get_appsettings,
    setup_logging,
)

from pyramid.scripts.common import parse_vars

from pyramid.scripts.common import parse_vars

from sqlalchemy import engine_from_config
from sqlalchemy.orm import sessionmaker

from ..models.meta import Base
from ..models import (
    get_session
)
from ..models import (Project,
                      User,
                      Equipment,
                      Obligation,
                      Member_table,
                      Member,
                      )

import datetime


def usage(argv):
    cmd = os.path.basename(argv[0])
    print('usage: %s <config_uri> [var=value]\n'
          '(example: "%s development.ini")' % (cmd, cmd))
    sys.exit(1)


def Get_data(session, project_id):
    return session.query(Project).filter_by(id=project_id).one()


def Gen_Doc_compet(doc_name='D.docx'
                   , project_name_th=u''
                   , project_name_en=u''
                   , date_cap=u''
                   , where=u''
                   , rational=u''
                   , purpose_list=list([])
                   , profit=u''
                   , owner_list=list([])
                   , advisor_list=u''
                   , member_list=list([])
                   , activity_place=u''
                   , type_of_activity=u''
                   , cost_list=list([])
                   , success_criteria=[]
                   ):
    document = Document('FRA241PROJECT/static/Archetype_gendoc.docx')

    ######################################################################
    '''Create Style'''
    ######################################################################

    obj_styles = document.styles

    obj_charstyle = obj_styles.add_style('header', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(16)
    obj_font.name = 'TH Sarabun New'
    obj_font.bold = True
    obj_font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    obj_charstyle = obj_styles.add_style('content', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(16)
    obj_font.name = 'TH Sarabun New'
    obj_font.bold = False
    obj_font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    obj_charstyle = obj_styles.add_style('in table', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(14)
    obj_font.name = 'TH Sarabun New'
    obj_font.bold = False
    obj_font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    ######################################################################

    document.add_picture('FRA241PROJECT/static/kmutt.png', width=Inches(0.99), height=Inches(0.99))
    logo = document.paragraphs[-1]
    logo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    proj_name = document.add_heading('', 0)
    proj_name.add_run(u'ข้อเสนอโครงการเข้าร่วม\n', style='header')
    proj_name.add_run(project_name_th + '\n', style='header') if len(project_name_th) != 0 else 0
    proj_name.add_run(project_name_en + '\n', style='header') if len(project_name_en) != 0 else 0
    proj_name.add_run(u'สถาบันวิทยาการหุ่นยนต์ภาคสนาม (ฟีโบ้) มหาวิทยาลัยเทคโนโลยีพระจอมเกล้าธนบุรี\n', style='header')
    proj_name.add_run(u'ระหว่างวันที่ ' + date_cap + '\n', style='header')
    proj_name.add_run(u'ณ ' + where, style='header')
    proj_name.alignment = WD_ALIGN_PARAGRAPH.CENTER

    rationale_head = document.add_paragraph('\n')
    rationale_head.add_run(u'หลักการและเหตุผล', style='header')

    rationale = document.add_paragraph('\t')
    rationale.add_run(rational + '\n', style='content')
    rationale.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY

    purpose_head = document.add_paragraph()
    purpose_head.add_run(u'วัตถุประสงค์', style='header').alignment = WD_ALIGN_PARAGRAPH.LEFT

    purpose = document.add_paragraph()
    x = 1
    for i in purpose_list:
        purpose.add_run('\t' + str(x) + '. ' + i + '\n', style='content') if len(i) != 0 else purpose.add_run()
        x += 1
    x = 0

    profit_head = document.add_paragraph()
    profit_head.add_run(u'ประโยชน์ที่คาดว่าจะได้รับ', style='header').alignment = WD_ALIGN_PARAGRAPH.LEFT

    profit_content = document.add_paragraph()
    x = 1
    for i in profit:
        profit_content.add_run('\t' + str(x) + '. ' + i + '\n', style='content') if len(
            i) != 0 else profit_content.add_run()
        x += 1
    x = 0

    owner_head = document.add_paragraph()
    owner_head.add_run(u'ผู้รับผิดชอบโครงการ', style='header').alignment = WD_ALIGN_PARAGRAPH.LEFT

    owner = document.add_paragraph()
    if len(owner_list) != 2:
        x = 1
        for i in owner_list:
            owner.add_run('\t' + str(x) + '. ' + i + '\n', style="content") if len(i) != 0 else owner.add_run()
            x += 1
        x = 0
    else:
        owner.add_run('\t' + owner_list[0] + '\n', style="content") if len(owner_list[0]) != 0 else owner.add_run()

    advisor_head = document.add_paragraph()
    advisor_head.add_run(u'อาจารย์ที่ปรึกษาโครงการ', style='header').alignment = WD_ALIGN_PARAGRAPH.LEFT

    advisor = document.add_paragraph()
    # if len(advisor_list) != 1 :
    #     x = 1
    #     for i in advisor_list:
    #         advisor.add_run('\t' + str(x) + '. ' + i + '\n', style="content")if len(i)!= 0 else advisor.add_run()
    #         x += 1
    #     x = 0
    # else:
    advisor.add_run('\t' + advisor_list + '\n', style="content") if len(advisor_list[0]) != 0 else advisor.add_run()

    member_head = document.add_paragraph()
    member_head.add_run(u'ผู้เข้าร่วมโครงการ', style='header').alignment = WD_ALIGN_PARAGRAPH.LEFT

    member = document.add_paragraph()
    x = 1
    for i in member_list:
        member.add_run('\t' + str(x) + '. ' + i + '\n', style="content") if len(i) != 0 else  member.add_run()
        x += 1
    x = 0

    plan_head = document.add_paragraph()
    plan_head.add_run(u'แผนการดำเนินงาน', style='header')

    plan = document.add_paragraph()
    plan.add_run(style='content')

    '''

    "Table"

        Date(details)

    '''

    place_head = document.add_paragraph()
    place_head.add_run(u'สถานที่ดำเนินการ', style='header').alignment = WD_ALIGN_PARAGRAPH.LEFT

    place = document.add_paragraph()
    place.add_run('\t' + activity_place, style='content')
    place.alingment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY

    activity_type_head = document.add_paragraph()
    activity_type_head.add_run(u'ลักษณะกิจกรรม', style='header').alignment = WD_ALIGN_PARAGRAPH.LEFT

    activity_type = document.add_paragraph()
    activity_type.add_run('\t' + type_of_activity, style='content')

    table_cost_head = document.add_paragraph()
    table_cost_head.add_run(u'ตารางแสดงค่าใช้จ่าย\n', style='header')

    table_cost = document.add_table(rows=1, cols=3)
    head_table = table_cost.rows[0].cells
    item_in_head_cost_table = [u'ลำดับ', u'รายละเอียด', u'จำนวนเงิน(บาท)']
    for i in range(0, 3):
        head_table_for = head_table[i].add_paragraph()
        head_table_for.add_run(item_in_head_cost_table[i], style='in table')

    for i in cost_list:
        each_cost = table_cost.add_row().cells
        for j in range(0, 3):
            each_cells = each_cost[j].add_paragraph()
            each_cells.add_run(i[j], style='in table')

    success_pointer_head = document.add_paragraph()
    success_pointer_head.add_run(u'\nตัวชี้วัดความสำเร็จของโครงการ', style='header')

    success_pointer = document.add_paragraph()
    if len(success_criteria) != 2:
        x = 1
        for i in success_criteria:
            success_pointer.add_run('\t' + str(x) + '. ' + i + '\n', style="content") if len(
                i) != 0 else success_pointer.add_run()
            x += 1
        x = 0
    else:
        success_pointer.add_run('\t' + success_criteria[0] + '\n', style="content") if len(
            success_criteria[0]) != 0 else success_pointer.add_run()

    sign_area = document.add_paragraph()
    sign_area.add_run(u'ลงชื่อ' + '.......................................................\n', style='content')
    sign_area.add_run((advisor_list.split('\t\t'))[0] + '\n', style="content")
    sign_area.add_run(u'อาจารย์ที่ปรึกษาโครงการ', style="content")
    sign_area.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # document.add_page_break()
    document.save(doc_name)

    print '#########################__Done!__#########################'


def Gen_Doc_camp(doc_name='D.docx'
                 , camp_name_th=u''
                 , project_name_en=u''
                 , year=u''
                 , date_cap=u''
                 , where=u''
                 , rational=u''
                 , purpose_list=list([])
                 , hours_compare=u''
                 , owner_list=list([])
                 , durations=u''
                 , member_list=list([])
                 , evaluation_index=u''
                 , profit=list([])
                 , cost_list=list([])
                 , cost_list_detail=list([])
                 , activity_list=list([])
                 ):
    document = Document('FRA241PROJECT/static/Archetype_gendoc.docx')

    ######################################################################
    '''Create Style'''
    ######################################################################

    obj_styles = document.styles

    obj_charstyle = obj_styles.add_style('header', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(16)
    obj_font.name = 'TH Sarabun New'
    obj_font.bold = True
    obj_font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    obj_charstyle = obj_styles.add_style('content', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(16)
    obj_font.name = 'TH Sarabun New'
    obj_font.bold = False
    obj_font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    obj_charstyle = obj_styles.add_style('in table', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(14)
    obj_font.name = 'TH Sarabun New'
    obj_font.bold = False
    obj_font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    ######################################################################

    document.add_picture('FRA241PROJECT/static/kmutt.png', width=Inches(0.99), height=Inches(0.99))
    logo = document.paragraphs[-1]
    logo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    head_camp_proposal = document.add_heading('', 0)
    head_camp_proposal.add_run(camp_name_th + '\n', style='header') if len(camp_name_th) != 0 else 0
    head_camp_proposal.add_run(project_name_en + '\n', style='header') if len(project_name_en) != 0 else 0
    head_camp_proposal.add_run(u'สถาบันวิทยาการหุ่นยนต์ภาคสนาม (ฟีโบ้) มหาวิทยาลัยเทคโนโลยีพระจอมเกล้าธนบุรี\n',
                               style='header')
    head_camp_proposal.add_run(u'ประจำปีการศึกษา ' + year + '\n', style='header')
    head_camp_proposal.add_run(u'ระหว่างวันที่ ' + date_cap + '\n', style='header')
    head_camp_proposal.add_run(u'ณ ' + where, style='header')
    head_camp_proposal.alignment = WD_ALIGN_PARAGRAPH.CENTER

    rationale_head = document.add_paragraph('\n')
    rationale_head.add_run(u'หลักการและเหตุผล', style='header')

    rationale = document.add_paragraph('\t')
    rationale.add_run(rational + '\n', style='content')

    purpose_head = document.add_paragraph()
    purpose_head.add_run(u'วัตถุประสงค์', style='header').alignment = WD_ALIGN_PARAGRAPH.LEFT

    purpose = document.add_paragraph()
    x = 1
    for i in purpose_list:
        purpose.add_run('\t' + str(x) + '. ' + i + '\n', style='content') if len(i) != 0 else purpose.add_run()
        x += 1
    x = 0

    '''
    การเทียบค่ากิจกรรม
    '''

    owner_head = document.add_paragraph()
    owner_head.add_run(u'ผู้รับผิดชอบโครงการ', style='header').alignment = WD_ALIGN_PARAGRAPH.LEFT

    owner = document.add_paragraph()
    if len(owner_list) != 2:
        x = 1
        for i in owner_list:
            owner.add_run('\t' + str(x) + '. ' + i + '\n', style="content") if len(i) != 0 else owner.add_run()
            x += 1
        x = 0
    else:
        owner.add_run('\t' + owner_list[0] + '\n', style="content") if len(owner_list[0]) != 0 else owner.add_run()

    duration_head = document.add_paragraph()
    duration_head.add_run(u'ระยะเวลาดำเนินงาน', style='header')

    duration = document.add_paragraph('\t')
    duration.add_run(durations + '\n', style='content')

    member_head = document.add_paragraph()
    member_head.add_run(u'ผู้เข้าร่วมโครงการ', style='header').alignment = WD_ALIGN_PARAGRAPH.LEFT

    member = document.add_paragraph()
    x = 1
    for i in member_list:
        member.add_run('\t' + str(x) + '. ' + i + '\n', style="content") if len(i) != 0 else  member.add_run()
        x += 1
    x = 0

    evaluation_head = document.add_paragraph()
    evaluation_head.add_run(u'รูปแบบการประเมิณผล', style='header')

    evaluation = document.add_paragraph('\t')
    evaluation.add_run(evaluation_index + '\n', style='content')

    profit_head = document.add_paragraph()
    profit_head.add_run(u'ผลที่คาดว่าจะได้รับ', style='header').alignment = WD_ALIGN_PARAGRAPH.LEFT

    profit_content = document.add_paragraph()
    x = 1
    for i in profit:
        profit_content.add_run('\t' + str(x) + '. ' + i + '\n', style='content') if len(
            i) != 0 else profit_content.add_run()
        x += 1
    x = 0

    cost_head = document.add_paragraph()
    cost_head.add_run(u'ค่าใช้จ่ายในการจัดอบรม', style='header').alignment = WD_ALIGN_PARAGRAPH.LEFT

    cost = document.add_paragraph()
    x = 1
    for i in cost_list:
        cost.add_run('\t' + str(x) + '. ' + i + '\n', style='content') if len(i) != 0 else profit_content.add_run()
        x += 1
    x = 0

    table_cost_head = document.add_paragraph()
    table_cost_head.add_run(u'งบประมาณที่ใช้มีรายละเอียดดังต่อไปนี้\n', style='header')

    table_cost = document.add_table(rows=1, cols=3)
    head_table = table_cost.rows[0].cells
    item_in_head_cost_table = [u'ลำดับที่', u'รายละเอียด', u'จำนวนเงิน(บาท)']
    for i in range(0, 3):
        head_table_for = head_table[i].add_paragraph()
        head_table_for.add_run(item_in_head_cost_table[i], style='in table')

    for i in cost_list_detail:
        each_cost = table_cost.add_row().cells
        for j in range(0, 3):
            each_cells = each_cost[j].add_paragraph()
            each_cells.add_run(i[j], style='in table')

    activity_table_head = document.add_paragraph()
    activity_table_head.add_run(u'\nตารางกิจกรรม', style='header').alignment = WD_ALIGN_PARAGRAPH.LEFT

    activity_table = document.add_paragraph()
    activity_double_list = []
    for i in activity_list:
        activity_double_list.append(i.split(unichr(172)))
    activity_double_list.pop()
    for j in activity_double_list:
        activity_table.add_run(j[0] + '\t' + j[1] + '\n', style='content')

    document.save(doc_name)


def Gen_Doc_volun(doc_name='D.docx'
                  , volun_name=u''
                  , year=u''
                  , where=u''
                  , rational=u''
                  , purpose_list=list([])
                  , hours_compare=u''
                  , advisor_list=list([])
                  , owner_list=list([])
                  , member_list=list([])
                  , durations=u''
                  , active_location=u''
                  , type_of_activity=u''
                  , evaluation_index=u''
                  , success_criteria=list([])
                  , profit=list([])
                  , cost_list_detail=list([])
                  ):
    document = Document('FRA241PROJECT/static/Archetype_gendoc.docx')

    ######################################################################
    '''Create Style'''
    ######################################################################

    obj_styles = document.styles

    obj_charstyle = obj_styles.add_style('header', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(16)
    obj_font.name = 'TH Sarabun New'
    obj_font.bold = True
    obj_font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    obj_charstyle = obj_styles.add_style('content', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(16)
    obj_font.name = 'TH Sarabun New'
    obj_font.bold = False
    obj_font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    obj_charstyle = obj_styles.add_style('in table', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(14)
    obj_font.name = 'TH Sarabun New'
    obj_font.bold = False
    obj_font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    ######################################################################

    document.add_picture('FRA241PROJECT/static/kmutt.png', width=Inches(0.99), height=Inches(0.99))
    logo = document.paragraphs[-1]
    logo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    head_volunteer_proposal = document.add_heading('', 0)
    head_volunteer_proposal.add_run(volun_name + '\n', style='header')
    head_volunteer_proposal.add_run(u'สาขาวิชาวิศวกรรมหุ่นยนต์และระบบอัตโนมัติ\n', style='header')
    head_volunteer_proposal.add_run(u'สถาบันวิทยาการหุ่นยนต์ภาคสนาม (ฟีโบ้) มหาวิทยาลัยเทคโนโลยีพระจอมเกล้าธนบุรี\n',
                                    style='header')
    head_volunteer_proposal.add_run(u'ประจำปีการศึกษา ' + year + '\n\n', style='header')
    head_volunteer_proposal.alignment = WD_ALIGN_PARAGRAPH.CENTER

    rationale_head = document.add_paragraph('\n')
    rationale_head.add_run(u'หลักการและเหตุผล', style='header').alignment = WD_ALIGN_PARAGRAPH.LEFT

    rationale = document.add_paragraph('\t')
    rationale.add_run(rational + '\n', style='content')

    purpose_head = document.add_paragraph()
    purpose_head.add_run(u'วัตถุประสงค์', style='header').alignment = WD_ALIGN_PARAGRAPH.LEFT

    purpose = document.add_paragraph()
    x = 1
    for i in purpose_list:
        purpose.add_run('\t' + str(x) + '. ' + i + '\n', style='content') if len(i) != 0 else purpose.add_run()
        x += 1
    x = 0

    advisor_head = document.add_paragraph()
    advisor_head.add_run(u'อาจารย์ที่ปรึกษาโครงการ', style='header').alignment = WD_ALIGN_PARAGRAPH.LEFT

    advisor = document.add_paragraph()
    x = 1
    for i in advisor_list:
        advisor.add_run('\t' + str(x) + '. ' + i + '\n', style='content') if len(i) != 0 else purpose.add_run()
        x += 1
    x = 0

    owner_head = document.add_paragraph()
    owner_head.add_run(u'ผู้รับผิดชอบโครงการ', style='header').alignment = WD_ALIGN_PARAGRAPH.LEFT

    owner = document.add_paragraph()
    if len(owner_list) != 2:
        x = 1
        for i in owner_list:
            owner.add_run('\t' + str(x) + '. ' + i + '\n', style="content") if len(i) != 0 else owner.add_run()
            x += 1
        x = 0
    else:
        owner.add_run('\t' + owner_list[0] + '\n', style="content") if len(owner_list[0]) != 0 else owner.add_run()

    member_head = document.add_paragraph()
    member_head.add_run(u'ผู้เข้าร่วมโครงการ', style='header').alignment = WD_ALIGN_PARAGRAPH.LEFT

    member = document.add_paragraph()
    x = 1
    for i in member_list:
        member.add_run('\t' + str(x) + '. ' + i + '\n', style="content") if len(i) != 0 else  member.add_run()
        x += 1
    x = 0

    duration_head = document.add_paragraph()
    duration_head.add_run(u'ระยะเวลาดำเนินงาน', style='header')

    duration = document.add_paragraph('\t')
    duration.add_run(durations + '\n', style='content')

    location_head = document.add_paragraph()
    location_head.add_run(u'สถานที่ปฏิบัติงาน', style='header').alignment = WD_ALIGN_PARAGRAPH.LEFT

    location = document.add_paragraph('\t')
    location.add_run(active_location + '\n', style='content')

    activity_type_head = document.add_paragraph()
    activity_type_head.add_run(u'ลักษณะกิจกรรม', style='header').alignment = WD_ALIGN_PARAGRAPH.LEFT

    activity_type = document.add_paragraph()
    activity_type.add_run('\t' + type_of_activity, style='content')

    evaluation_head = document.add_paragraph()
    evaluation_head.add_run(u'รูปแบบการประเมิณผล', style='header')

    evaluation = document.add_paragraph('\t')
    evaluation.add_run(evaluation_index + '\n', style='content')

    success_pointer_head = document.add_paragraph()
    success_pointer_head.add_run(u'\nตัวชี้วัดความสำเร็จของโครงการ', style='header')

    success_pointer = document.add_paragraph()
    if len(success_criteria) != 2:
        x = 1
        for i in success_criteria:
            success_pointer.add_run('\t' + str(x) + '. ' + i + '\n', style="content") if len(
                i) != 0 else success_pointer.add_run()
            x += 1
        x = 0
    else:
        success_pointer.add_run('\t' + success_criteria[0] + '\n', style="content") if len(
            success_criteria[0]) != 0 else success_pointer.add_run()

    profit_head = document.add_paragraph()
    profit_head.add_run(u'ผลที่คาดว่าจะได้รับ', style='header').alignment = WD_ALIGN_PARAGRAPH.LEFT

    profit_content = document.add_paragraph()
    x = 1
    for i in profit:
        profit_content.add_run('\t' + str(x) + '. ' + i + '\n', style='content') if len(
            i) != 0 else profit_content.add_run()
        x += 1
    x = 0

    table_cost_head = document.add_paragraph()
    table_cost_head.add_run(u'งบประมาณที่ใช้\n', style='header')
    table_cost_head.add_run(u'งบประมาณที่ใช้มีรายละเอียดดังต่อไปนี้\n', style='content')
    table_cost_head.alignment = WD_ALIGN_PARAGRAPH.LEFT

    table_cost = document.add_table(rows=1, cols=3)
    head_table = table_cost.rows[0].cells
    item_in_head_cost_table = [u'ลำดับที่', u'รายละเอียด', u'จำนวนเงิน(บาท)']
    for i in range(0, 3):
        head_table_for = head_table[i].add_paragraph()
        head_table_for.add_run(item_in_head_cost_table[i], style='in table')

    for i in cost_list_detail:
        each_cost = table_cost.add_row().cells
        for j in range(0, 3):
            each_cells = each_cost[j].add_paragraph()
            each_cells.add_run(i[j], style='in table')

    document.save(doc_name)


def gennn(ses, num, path):
    month_list = [u'มกราคม', u'กุมภาพันธ์', u'มีนาคม', u'เมษายน', u'พฤษภาคม', u'มิถุนายน', u'กรกฏาคม', u'สิงหาคม',
                  u'กันยายน', u'ตุลาคม', u'พฤศจิกายน', u'ธันวาคม']

    # all_data = Get_data(session=session, project_id=int(argv[2]))

    all_data = Get_data(session=ses, project_id=num)

    cost_list_parameter = []
    cost_list_param = all_data.proposal.delicate_budget.split(
        unichr(171)) if all_data.proposal.cost is not None else []
    cost_list_param.pop()

    for i in cost_list_param:
        cost_list_parameter.append(i.split(unichr(172)))

    if all_data.type == u'competitive':
        startdate = all_data.start_date
        findate = all_data.finish_date
        if startdate is not None:
            datecap = unicode(
                unicode(startdate.day) + ' ' + month_list[int(startdate.month) + 1] + ' - ' + unicode(
                    startdate.day) + ' ' +
                month_list[int(findate.month) + 1] + ' ' + unicode(findate.year))
        else:
            datecap = u''
        Gen_Doc_compet(doc_name='FRA241PROJECT/static/Gened_DOC/' + 'competitive_' + str(all_data.id) + '.docx'
                       , project_name_th=all_data.title
                       , date_cap=datecap
                       , where=all_data.proposal.activity_location
                       , rational=all_data.proposal.Reason
                       , purpose_list=all_data.proposal.objective.split(unichr(171))
                       , profit=all_data.proposal.profit.split(unichr(171))
                       , owner_list=all_data.proposal.owner_for_proposal.split(unichr(171))
                       , advisor_list=all_data.proposal.advisor_for_proposal
                       , member_list=all_data.proposal.member_for_proposal.split(unichr(171))
                       , activity_place=all_data.proposal.activity_location
                       , type_of_activity=all_data.proposal.type_of_activity
                       , cost_list=cost_list_parameter
                       , success_criteria=all_data.proposal.success_criteria.split(unichr(171))
                       )
        # Gen_Doc_compet(doc_name=path + 'Competitive_' + str(all_data.id) + '.text'
        #                , project_name_th=all_data.title
        #                , date_cap=u''
        #                , where=all_data.proposal.activity_location
        #                , rational=all_data.proposal.Reason
        #                , purpose_list=all_data.proposal.objective.split(unichr(171))
        #                , profit=all_data.proposal.profit.split(unichr(171))
        #                , owner_list=all_data.proposal.owner_for_proposal.split(unichr(171))
        #                , advisor_list=all_data.proposal.advisor_for_proposal
        #                , member_list=all_data.proposal.member_for_proposal.split(unichr(171))
        #                , activity_place=all_data.proposal.activity_location
        #                , type_of_activity=all_data.proposal.type_of_activity
        #                , cost_list=cost_list_parameter
        #                , success_criteria=all_data.proposal.success_criteria.split(unichr(171))
        #                )
        # os.startfile('C:\Users\PHURINPAT\Documents\GitHub\FRA241Group6\FRA241PROJECT\static\Gened_DOC\Competitive_2.docx')
    elif all_data.type == u'camp':
        startdate = all_data.start_date
        findate = all_data.finish_date
        if startdate is not None:
            datecap = unicode(
                unicode(startdate.day) + ' ' + month_list[int(startdate.month) + 1] + ' - ' + unicode(
                    startdate.day) + ' ' +
                month_list[int(findate.month) + 1] + ' ' + unicode(findate.year))
        else:
            datecap = u''

        Gen_Doc_camp(doc_name='FRA241PROJECT/static/Gened_DOC/' + 'camp_' + str(all_data.id) + '.docx'
                     , camp_name_th=all_data.title
                     , year=all_data.proposal.year
                     , date_cap=datecap
                     , where=all_data.proposal.activity_location
                     , rational=all_data.proposal.Reason
                     , purpose_list=all_data.proposal.objective.split(unichr(171))
                     , hours_compare=u''
                     , owner_list=all_data.proposal.owner_for_proposal.split(unichr(171))
                     , durations=all_data.proposal.duration
                     , member_list=all_data.proposal.member_for_proposal.split(unichr(171))
                     , evaluation_index=all_data.proposal.evaluation_index
                     , profit=all_data.proposal.profit.split(unichr(171))
                     , cost_list=all_data.proposal.cost.split(unichr(171))
                     , cost_list_detail=cost_list_parameter
                     , activity_list=all_data.proposal.schedule.split(unichr(171))
                     )
        # Gen_Doc_camp(doc_name=path + 'Camp_' + str(all_data.id) + '.text'
        #              , camp_name_th=all_data.title
        #              , year=all_data.proposal.year
        #              , date_cap=u''
        #              , where=all_data.proposal.activity_location
        #              , rational=all_data.proposal.Reason
        #              , purpose_list=all_data.proposal.objective.split(unichr(171))
        #              , hours_compare=u''
        #              , owner_list=all_data.proposal.owner_for_proposal.split(unichr(171))
        #              , durations=all_data.proposal.duration
        #              , member_list=all_data.proposal.member_for_proposal.split(unichr(171))
        #              , evaluation_index=all_data.proposal.evaluation_index
        #              , profit=all_data.proposal.profit.split(unichr(171))
        #              , cost_list=all_data.proposal.cost.split(unichr(171))
        #              , cost_list_detail=cost_list_parameter
        #              , activity_list=all_data.proposal.schedule.split(unichr(171))
        #              )
        # os.startfile('C:\Users\PHURINPAT\Documents\GitHub\FRA241Group6\sdf.docx')
    elif all_data.type == u'volunteer':
        Gen_Doc_volun(doc_name='FRA241PROJECT/static/Gened_DOC/' + 'volunteer_' + str(all_data.id) + '.docx'
                      , volun_name=all_data.title
                      , year=all_data.proposal.year
                      , where=all_data.proposal.activity_location
                      , rational=all_data.proposal.Reason
                      , purpose_list=all_data.proposal.objective.split(unichr(171))
                      , hours_compare=u''
                      , advisor_list=all_data.proposal.advisor_for_proposal.split(unichr(171))
                      , owner_list=all_data.proposal.owner_for_proposal.split(unichr(171))
                      , member_list=all_data.proposal.member_for_proposal.split(unichr(171))
                      , durations=all_data.proposal.duration
                      , active_location=all_data.proposal.activity_location
                      , type_of_activity=all_data.proposal.type_of_activity
                      , evaluation_index=all_data.proposal.evaluation_index
                      , success_criteria=all_data.proposal.success_criteria.split(unichr(171))
                      , profit=all_data.proposal.profit.split(unichr(171))
                      , cost_list_detail=cost_list_parameter
                      )
        # Gen_Doc_volun(doc_name=path + 'Volun_' + str(all_data.id) + '.text'
        #               , volun_name=all_data.title
        #               , year=all_data.proposal.year
        #               , where=all_data.proposal.activity_location
        #               , rational=all_data.proposal.Reason
        #               , purpose_list=all_data.proposal.objective.split(unichr(171))
        #               , hours_compare=u''
        #               , advisor_list=all_data.proposal.advisor_for_proposal.split(unichr(171))
        #               , owner_list=all_data.proposal.owner_for_proposal.split(unichr(171))
        #               , member_list=all_data.proposal.member_for_proposal.split(unichr(171))
        #               , durations=all_data.proposal.duration
        #               , active_location=all_data.proposal.activity_location
        #               , type_of_activity=all_data.proposal.type_of_activity
        #               , evaluation_index=all_data.proposal.evaluation_index
        #               , success_criteria=all_data.proposal.success_criteria.split(unichr(171))
        #               , profit=all_data.proposal.profit.split(unichr(171))
        #               , cost_list_detail=cost_list_parameter
        #               )


def main(argv=sys.argv):
    if len(argv) < 2:
        usage(argv)
    config_uri = argv[1]
    options = parse_vars(argv[3:])
    setup_logging(config_uri)
    settings = get_appsettings(config_uri, options=options)

    engine = engine_from_config(settings, prefix='sqlalchemy.')
    Base.metadata.create_all(engine)

    maker = sessionmaker()
    maker.configure(bind=engine)

    session = get_session(maker, transaction.manager)

    gennn(session, int(argv[2]))
